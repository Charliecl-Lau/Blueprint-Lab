import hashlib
from io import BytesIO

import pytest
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from backend.models.source_document import SourceDocument
from backend.services.source_documents import SourceDocumentValidationError, create_source_document


def create_docx() -> bytes:
    stream = BytesIO()
    document = Document()
    document.add_paragraph("First paragraph")
    document.add_paragraph("Second paragraph")
    document.save(stream)
    return stream.getvalue()


def create_pdf() -> bytes:
    stream = BytesIO()
    writer = PdfWriter()
    page = writer.add_blank_page(width=200, height=200)
    content = DecodedStreamObject()
    content.set_data(b"BT /F1 12 Tf 20 100 Td (PDF page text) Tj ET")
    page[NameObject("/Contents")] = writer._add_object(content)
    page[NameObject("/Resources")] = DictionaryObject({NameObject("/Font"): DictionaryObject({NameObject("/F1"): writer._add_object(
        DictionaryObject({
            NameObject("/Type"): NameObject("/Font"), NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica")
        })
    )})})
    writer.write(stream)
    return stream.getvalue()


def create(db, content=b"hello \xe2\x98\x83", media_type="text/plain", filename="source.txt"):
    return create_source_document(db, name="Syllabus", document_type="course_syllabus",
        version="2026", filename=filename, media_type=media_type, content=content, description="Current")


def test_text_snapshot_keeps_exact_bytes_and_hash(test_db):
    content = b"hello \xe2\x98\x83"
    source = create(test_db, content)
    assert source.content == content
    assert source.content_hash == hashlib.sha256(content).hexdigest()
    assert source.extracted_text == "hello \u2603"
    assert source.extraction_method == "plain-text:utf-8"


@pytest.mark.parametrize(("content", "media_type", "filename", "text", "method"), [
    (create_docx(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "a.docx", "First paragraph\nSecond paragraph", "python-docx:"),
    (create_pdf(), "application/pdf", "a.pdf", "PDF page text", "pypdf:"),
], ids=["docx", "pdf"])
def test_binary_snapshot_extracts_text_without_altering_bytes(test_db, content, media_type, filename, text, method):
    source = create(test_db, content, media_type, filename)
    assert source.content == content
    assert text in source.extracted_text
    assert source.extraction_method.startswith(method)


@pytest.mark.parametrize(("content", "media_type", "code"), [
    (b"x" * (20 * 1024 * 1024 + 1), "text/plain", "source_document_too_large"),
    (b"hello", "image/png", "unsupported_source_document_media_type"),
    (b"\xff", "text/plain", "invalid_source_document"),
], ids=["too-large", "unsupported", "invalid-utf8"])
def test_invalid_upload_does_not_insert(test_db, content, media_type, code):
    with pytest.raises(SourceDocumentValidationError) as exc:
        create(test_db, content, media_type)
    assert exc.value.code == code
    assert test_db.query(SourceDocument).count() == 0
