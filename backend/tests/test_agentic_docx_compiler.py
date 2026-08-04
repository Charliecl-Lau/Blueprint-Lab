import io
import zipfile

from docx import Document
from docx.oxml.ns import qn

from backend.schemas.docx_tool_schema import DocxToolCall
from backend.services.agentic_docx_compiler import AgenticDocxCompiler
from backend.services.docx_content_catalog import DocxContentCatalog
from backend.services.docx_tool_workspace import DocxWorkspace


def _call(index, tool, **args):
    return DocxToolCall(operation_id=f"op-{index}", tool=tool, expected_workspace_revision=index, arguments=args)


def test_compiler_clones_assessment_and_emits_native_math_and_evidence():
    source = {"questions": [{
        "type": "short_answer", "id": "1", "metadata": {"question_title": "Energy"},
        "body": "Use [[EQ:e]].", "options": [], "model_answer": "It is [[EQ:a]].",
        "equations": [
            {"label": "e", "math": {"type": "symbol", "name": "E"}, "location": "question"},
            {"label": "a", "math": {"type": "symbol", "name": "A"}, "location": "solution"},
        ],
        "quality_checks": [{"criterion": "Correctness", "rating": 5, "comment": "Correct."}],
        "revision_options": ["Change material", "Change temperature"],
    }]}
    catalog = DocxContentCatalog.from_assessment(source)
    workspace = DocxWorkspace.create(catalog)
    workspace.apply_batch([
        _call(0, "create_document"),
        _call(1, "add_section", block_id="questions", role="questions"),
        _call(2, "add_question", block_id="question-1", parent_id="questions", question_id="1"),
        _call(3, "add_section", block_id="solutions", role="solutions"),
        _call(4, "add_solution", block_id="solution-1", parent_id="solutions", question_id="1"),
        _call(5, "finalize_document"),
    ])
    result = AgenticDocxCompiler().compile(workspace, session_id=3, iteration_number=0)
    assert result.assessment_json == source
    assert result.layout_manifest["docx_hash"] == result.docx_sha256
    assert len(result.layout_manifest["equation_placements"]) == 2
    with zipfile.ZipFile(io.BytesIO(result.docx_bytes)) as archive:
        xml = archive.read("word/document.xml")
    assert b"<m:oMath" in xml
    assert "Question 1" in "\n".join(p.text for p in Document(io.BytesIO(result.docx_bytes)).paragraphs)


def test_compiler_applies_reference_question_solution_and_table_style():
    source = {
        "metadata": {"course": "MSE302 Thermodynamics II", "topic": "Solutions"},
        "questions": [{
            "type": "mcq", "id": "1",
            "metadata": {
                "question_title": "Chemical potential from activity",
                "question_type": "Multiple choice",
                "difficulty_level": "Medium",
                "intended_assessment_setting": "Quiz",
                "mse202_concepts": ["Chemical potential"],
                "mse302_concepts": ["Activity"],
                "concept_map_bridge": "Extends unary chemical potential to components.",
                "materials_science_context": "Binary alloy",
                "estimated_time": "5 minutes",
                "learning_objectives": ["Calculate chemical potential"],
            },
            "body": "Choose the best expression.",
            "options": [
                {"body": "First", "is_correct": True},
                {"body": "Second", "is_correct": False},
                {"body": "Third", "is_correct": False},
                {"body": "Fourth", "is_correct": False},
                {"body": "Fifth", "is_correct": False},
            ],
            "model_answer": (
                "Step 1 — Identify the governing relation\n"
                "[[EQ:governing]]\n"
                "This definition applies to the stated system.\n\n"
                "Step 2 — Substitute the known values\n"
                "[[EQ:substitution]]\n"
                "The units remain consistent.\n\n"
                "Step 3 — State and interpret the answer\n"
                "The first choice is correct and has the expected physical sign.\n\n"
                "Why the other choices are incorrect\n"
                "B. It uses the wrong sign.\n"
                "C. It omits the activity term.\n"
                "D. It has inconsistent units.\n"
                "E. It applies the wrong reference state."
            ),
            "equations": [
                {"label": "governing", "expression": "mu_i = mu_i^o + R T ln(a_i)", "location": "solution"},
                {"label": "substitution", "expression": "mu_i = -20.0 kJ/mol + (8.314e-3 kJ/(mol K))(1200 K) ln(0.300)", "location": "solution"},
            ],
            "revision_options": ["Increase the numerical difficulty."],
            "quality_checks": [{"criterion": "Technical correctness", "rating": 5, "comment": "Correct."}],
        }],
    }
    catalog = DocxContentCatalog.from_assessment(source)
    workspace = DocxWorkspace.create(catalog)
    operations = [
        ("create_document", {}),
        ("add_section", {"block_id": "metadata", "role": "assessment_metadata"}),
        ("add_section", {"block_id": "questions", "role": "questions"}),
        ("add_question", {"block_id": "q1", "parent_id": "questions", "question_id": "1"}),
        ("add_section", {"block_id": "solutions", "role": "solutions"}),
        ("add_solution", {"block_id": "s1", "parent_id": "solutions", "question_id": "1"}),
        ("add_section", {"block_id": "key", "role": "answer_key"}),
        ("add_answer_key", {"block_id": "key-table", "parent_id": "key"}),
        ("add_section", {"block_id": "quality", "role": "quality_check"}),
        ("add_quality_check", {"block_id": "quality-table", "parent_id": "quality"}),
        ("add_section", {"block_id": "revisions", "role": "revision_options"}),
        ("add_content", {"block_id": "revision-1", "parent_id": "revisions", "content_ref": "question.1.revision_option.0"}),
        ("finalize_document", {}),
    ]
    workspace.apply_batch([_call(index, tool, **arguments) for index, (tool, arguments) in enumerate(operations)])

    result = AgenticDocxCompiler().compile(workspace)
    document = Document(io.BytesIO(result.docx_bytes))
    text = [paragraph.text for paragraph in document.paragraphs]

    assert document.styles["Normal"].font.name == "Aptos"
    assert document.styles["Normal"].font.size.pt == 10
    assert round(document.sections[0].left_margin.inches, 2) == 0.72
    assert document.sections[0].header.paragraphs[0].text == (
        "MSE302 Thermodynamics II  |  Blueprint Solutions Question Bank"
    )
    assert "Question 1 — Chemical potential from activity" in text
    assert "Solution 1 — Chemical potential from activity" in text
    assert "A. First" in text and "E. Fifth" in text
    assert "Correct answer: A." in text
    assert text.index("Answer key") < text.index("Solution 1 — Chemical potential from activity")
    assert not any(paragraph.text.startswith("Step ") for paragraph in document.paragraphs)
    assert [paragraph.text for paragraph in document.paragraphs if paragraph.style.name == "Solution Subheading"] == [
        "Why the other choices are incorrect",
    ]
    assert len([paragraph for paragraph in document.paragraphs if paragraph.style.name == "Solution Equation"]) == 2
    assert next(paragraph for paragraph in document.paragraphs if paragraph.text.startswith("B.")).runs[0].bold
    assert [len(table.columns) for table in document.tables] == [2, 2, 5]
    assert document.tables[0].cell(0, 0)._tc.tcPr.find(qn("w:shd")).get(qn("w:fill")) == "1F4E79"
    assert document.tables[2].cell(0, 0).text == "Criterion"


def test_compiler_ignores_model_architecture_and_uses_fixed_template():
    source = {"questions": [
        {
            "type": "short_answer",
            "id": "2",
            "metadata": {"question_title": "Second"},
            "body": "Second body.",
            "options": [],
            "model_answer": "Second solution.",
            "equations": [],
            "quality_checks": [{"criterion": "Clarity", "rating": 5, "comment": "Clear."}],
            "revision_options": ["Revise second."],
        },
        {
            "type": "short_answer",
            "id": "1",
            "metadata": {"question_title": "First"},
            "body": "First body.",
            "options": [],
            "model_answer": "First solution.",
            "equations": [],
            "quality_checks": [{"criterion": "Clarity", "rating": 5, "comment": "Clear."}],
            "revision_options": ["Revise first."],
        },
    ]}
    catalog = DocxContentCatalog.from_assessment(source)
    workspace = DocxWorkspace.create(catalog)
    operations = [
        ("create_document", {}),
        ("add_heading", {"block_id": "generic-questions", "literal_text": "Questions"}),
        ("add_section", {"block_id": "solutions", "role": "solutions"}),
        ("add_solution", {"block_id": "s1", "parent_id": "solutions", "question_id": "1"}),
        ("add_solution", {"block_id": "s2", "parent_id": "solutions", "question_id": "2"}),
        ("add_section", {"block_id": "questions", "role": "questions"}),
        ("add_question", {"block_id": "q1", "parent_id": "questions", "question_id": "1"}),
        ("add_question", {"block_id": "q2", "parent_id": "questions", "question_id": "2"}),
        ("add_heading", {"block_id": "generic-solutions", "literal_text": "Solutions"}),
        ("finalize_document", {}),
    ]
    workspace.apply_batch([
        _call(index, tool, **arguments)
        for index, (tool, arguments) in enumerate(operations)
    ])

    result = AgenticDocxCompiler().compile(workspace)
    text = [p.text for p in Document(io.BytesIO(result.docx_bytes)).paragraphs]
    sections = [
        "1. Assessment Metadata",
        "2. Student-Facing Questions",
        "3. Fully Worked Solution",
        "4. Assessment Quality Check",
        "5. Suggested Revision Options",
    ]
    assert [text.index(section) for section in sections] == sorted(
        text.index(section) for section in sections
    )
    assert text.index("Question 1 — Second") < text.index("Question 2 — First")
    assert text.index("Question 2 — First") < text.index("Solution 1 — Second")
    assert "Questions" not in text and "Solutions" not in text
