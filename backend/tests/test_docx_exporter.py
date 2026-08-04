from io import BytesIO
from copy import deepcopy
from zipfile import ZipFile

import pytest
from docx import Document
from lxml import etree

from backend.services.docx_exporter import (
    DocxExportValidationError,
    build_assessment_docx as _build_assessment_docx,
)


def build_assessment_docx(**kwargs):
    """Keep focused rendering fixtures valid under the export preflight contract."""
    kwargs = deepcopy(kwargs)
    for index, question in enumerate(kwargs.get("questions", []), start=1):
        question.setdefault("metadata", {})
        question["metadata"].setdefault("question_title", f"Fixture {index}")
        question.setdefault("quality_checks", [{
            "criterion": "Technical correctness",
            "rating": 5,
            "comment": "Verified for this fixture.",
        }])
        question.setdefault("revision_options", ["Vary the supplied values."])
        if not question["revision_options"]:
            question["revision_options"] = ["Vary the supplied values."]
    return _build_assessment_docx(**kwargs)


def thermodynamic_equation_ast():
    return {
        "type": "equation",
        "left": {
            "type": "fraction",
            "numerator": {"type": "differential", "variable": "P"},
            "denominator": {"type": "differential", "variable": "T"},
        },
        "right": {
            "type": "fraction",
            "numerator": {"type": "symbol", "name": "DeltaH"},
            "denominator": {
                "type": "product",
                "terms": [
                    {"type": "symbol", "name": "T"},
                    {"type": "symbol", "name": "DeltaV"},
                ],
            },
        },
    }


def table_rows(document):
    return [
        tuple(cell.text for cell in row.cells)
        for table in document.tables
        for row in table.rows
    ]


def test_docx_displays_application_traceability_and_omits_removed_setting():
    content = build_assessment_docx(
        run_id=12,
        prompt_id=34,
        condition_code="C101",
        run_number=2,
        course="MSE302",
        topic="Phase equilibrium",
        traceability={
            "experiment_id": 5,
            "condition_id": 7,
            "run_id": 12,
            "prompt_id": 34,
            "assessment_id": 56,
        },
        questions=[{
            "traceability": {
                "assessment_question_id": 78,
                "ordinal": 0,
                "assessment_version": 1,
            },
            "metadata": {
                "question_title": "Phase stability",
                "question_type": "short_answer",
                "difficulty_level": "medium",
                "mse202_concepts": ["Equilibrium"],
                "mse302_concepts": ["Gibbs energy"],
                "concept_map_bridge": None,
                "materials_science_context": "Binary alloy",
                "estimated_time_minutes": 10,
                "learning_objectives": ["Compare Gibbs energies"],
            },
            "body": "Explain stability.",
            "options": [],
            "model_answer": "Compare Gibbs energies.",
            "equations": [],
            "revision_options": [],
        }],
    )

    document = Document(BytesIO(content))
    rows = table_rows(document)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert ("Experiment ID", "5") in rows
    assert ("Condition ID", "7") in rows
    assert ("Assessment ID", "56") in rows
    assert "Question ID 78" in text
    assert "Intended Assessment Setting" not in str(rows)


def test_docx_uses_one_metadata_table_sourced_from_first_question():
    content = build_assessment_docx(
        run_id=12,
        prompt_id=34,
        condition_code="C101",
        run_number=2,
        course="MSE302",
        topic="Phase equilibrium",
        questions=[
            {
                "metadata": {
                    "question_title": "Shared equilibrium assessment",
                    "question_type": "long_answer",
                    "difficulty_level": "Medium",
                    "mse202_concepts": ["Gibbs Phase Rule"],
                },
                "body": "Explain the single-phase region.",
                "options": [],
                "model_answer": "It has two degrees of freedom.",
                "revision_options": [],
            },
            {
                "metadata": {
                    "question_title": "This title must not become metadata",
                    "difficulty_level": "Hard",
                },
                "body": "Explain the eutectic point.",
                "options": [],
                "model_answer": "It is invariant at fixed pressure.",
                "revision_options": [],
            },
        ],
    )

    document = Document(BytesIO(content))
    rows = table_rows(document)
    assert len(document.tables) == 2
    assert document.tables[0].cell(0, 0).text == "Field"
    assert document.tables[0].cell(0, 1).text == "Entry"
    assert ("Run ID", "12") in rows
    assert ("Prompt ID", "34") in rows
    assert ("Condition Code", "C101") in rows
    assert ("Run Number", "2") in rows
    assert ("Course", "MSE302") in rows
    assert ("Topic", "Phase equilibrium") in rows
    assert ("Question Title", "Shared equilibrium assessment") in rows
    assert ("Difficulty Level", "Medium") in rows
    assert ("MSE202 Concept(s)", "Gibbs Phase Rule") in rows
    assert all("This title must not become metadata" not in value for row in rows for value in row)
    assert all(value != "Hard" for row in rows for value in row)


def test_docx_empty_assessment_fails_preflight_visibly():
    with pytest.raises(DocxExportValidationError, match="questions are required"):
        _build_assessment_docx(
            run_id=1,
            prompt_id=2,
            condition_code="C100",
            run_number=1,
            course="ENGR 101",
            topic="Statics",
            questions=[],
        )


def test_docx_applies_spaced_item_styles_and_real_answer_choice_lists():
    content = build_assessment_docx(
        run_id=3,
        prompt_id=4,
        condition_code="C001",
        run_number=1,
        course="MSE202",
        topic="Phase rule",
        questions=[{
            "metadata": {"question_title": "Phase count"},
            "body": "How many phases coexist?",
            "options": [
                {"body": "One", "is_correct": False},
                {"body": "Three", "is_correct": True},
            ],
            "model_answer": "Three phases coexist.",
            "revision_options": [],
        }],
    )

    document = Document(BytesIO(content))
    paragraphs = document.paragraphs
    assert any(p.text == "Question 1 — Phase count" and p.style.name == "Heading 2" for p in paragraphs)
    assert any(p.text == "Solution 1 — Phase count" and p.style.name == "Heading 2" for p in paragraphs)
    choices = [p for p in paragraphs if p.text in {"One", "Three"}]
    assert len(choices) == 2
    assert all(p.style.name == "List Bullet" for p in choices)
    assert all(not p.text.startswith("- ") for p in choices)
    assert document.styles["Heading 2"].paragraph_format.space_before.pt == 12
    assert document.styles["Heading 3"].paragraph_format.space_before.pt >= 12
    assert document.styles["Normal"].paragraph_format.space_after.pt >= 6
    assert document.styles["List Bullet"].paragraph_format.space_after.pt >= 4


def test_docx_contains_rich_research_content_and_native_word_equation():
    content = build_assessment_docx(
        run_id=12, prompt_id=34, condition_code="C101", run_number=2,
        course="MSE302", topic="Phase equilibrium",
        questions=[{
            "type": "mcq",
            "metadata": {
                "question_title": "Equilibrium condition",
                "concept_map_bridge": "Connects MSE202 free energy to MSE302 phase stability.",
                "materials_science_context": "Determines stable phases in an alloy.",
            },
            "body": "Which condition identifies phase equilibrium?",
            "options": [
                {"body": "Equal chemical potentials", "is_correct": True},
                {"body": "Unequal temperatures", "is_correct": False},
            ],
            "model_answer": "Chemical potentials are equal at equilibrium.",
            "equations": [{"label": "Equilibrium", "expression": "mu_alpha = mu_beta", "location": "solution"}],
            "quality_checks": [{"criterion": "Correctness", "rating": 5, "comment": "Thermodynamically correct."}],
            "revision_options": ["Ask students to derive the equilibrium condition."],
        }],
    )

    document = Document(BytesIO(content))
    rows = table_rows(document)
    text = "\n".join(p.text for p in document.paragraphs)
    assert ("Run ID", "12") in rows
    assert ("Prompt ID", "34") in rows
    assert ("Condition Code", "C101") in rows
    assert ("Run Number", "2") in rows
    assert ("Concept-Map Bridge", "Connects MSE202 free energy to MSE302 phase stability.") in rows
    assert "Chemical potentials are equal at equilibrium." in text
    assert "4. Assessment Quality Check" in text
    assert "5. Suggested Revision Options" in text
    assert "End-to-end token usage" not in text

    with ZipFile(BytesIO(content)) as archive:
        document_xml = archive.read("word/document.xml")
    assert b"<m:oMath" in document_xml
    assert document_xml.count(b"<m:sSub>") == 2
    assert "α".encode() in document_xml
    assert "β".encode() in document_xml


def test_docx_converts_flat_word_linear_equations_to_built_up_omml():
    content = build_assessment_docx(
        run_id=20, prompt_id=21, condition_code="C010", run_number=1,
        course="MSE202", topic="Thermodynamics",
        questions=[{
            "metadata": {},
            "body": "Evaluate the expressions.",
            "options": [],
            "model_answer": "Use the equations shown.",
            "equations": [
                {
                    "label": "Fraction",
                    "expression": "DeltaH/(T DeltaS)",
                    "location": "solution",
                },
                {
                    "label": "Scripts",
                    "expression": "x_a^2",
                    "location": "solution",
                },
                {
                    "label": "Radical",
                    "expression": "sqrt(x_a)",
                    "location": "solution",
                },
            ],
        }],
    )

    with ZipFile(BytesIO(content)) as archive:
        document_xml = archive.read("word/document.xml")

    assert b"<m:f>" in document_xml
    assert b"<m:num>" in document_xml
    assert b"<m:den>" in document_xml
    assert document_xml.count(b"<m:sSub>") >= 2
    assert b"<m:sSup>" in document_xml
    assert b"<m:rad>" in document_xml


def test_docx_keeps_complete_signed_exponent_inside_omml_superscript():
    content = build_assessment_docx(
        run_id=30,
        prompt_id=31,
        condition_code="C001",
        run_number=1,
        course="MSE202",
        topic="Signed powers",
        questions=[{
            "metadata": {},
            "body": "Interpret the unit.",
            "options": [],
            "model_answer": "Use the unit shown.",
            "equations": [{
                "label": "InverseUnit",
                "expression": "K^-1",
                "location": "solution",
            }],
        }],
    )

    with ZipFile(BytesIO(content)) as archive:
        document_xml = archive.read("word/document.xml")

    root = etree.fromstring(document_xml)
    namespace = {
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    }
    superscripts = root.xpath("//m:sSup/m:sup", namespaces=namespace)

    assert len(superscripts) == 1
    assert "".join(superscripts[0].itertext()) == "-1"


def test_docx_replaces_equation_placeholders_inline_without_duplicate_blocks():
    content = build_assessment_docx(
        run_id=22, prompt_id=23, condition_code="C011", run_number=1,
        course="MSE202", topic="Thermodynamics",
        questions=[{
            "metadata": {},
            "body": "Use [[EQ:gibbs_formula]] to calculate the change.",
            "options": [{
                "body": "The value is [[EQ:option_value]].",
                "is_correct": True,
            }],
            "model_answer": "Substitution gives [[EQ:final_result]].",
            "equations": [
                {
                    "label": "gibbs_formula",
                    "expression": "DeltaG=DeltaH-T DeltaS",
                    "location": "question",
                },
                {
                    "label": "option_value",
                    "expression": "x_a^2",
                    "location": "question",
                },
                {
                    "label": "final_result",
                    "expression": "DeltaG=-180 J/mol",
                    "location": "solution",
                },
            ],
        }],
    )

    with ZipFile(BytesIO(content)) as archive:
        document_xml = archive.read("word/document.xml")

    assert b"[[EQ:" not in document_xml
    assert document_xml.count(b"<m:oMath>") == 3
    assert b"gibbs_formula:" not in document_xml
    assert b"option_value:" not in document_xml
    assert b"final_result:" not in document_xml
    assert b"<m:sSup>" in document_xml


def test_docx_omits_end_to_end_token_usage():
    content = build_assessment_docx(
        run_id=1,
        prompt_id=2,
        condition_code="C100",
        run_number=1,
        course="ENGR 101",
        topic="Statics",
        questions=[{
            "metadata": {"question_title": "Statics"},
            "body": "State equilibrium.",
            "options": [],
            "model_answer": "The net force and moment are zero.",
            "equations": [],
        }],
    )

    document = Document(BytesIO(content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "End-to-end token usage" not in text
    assert "Input:" not in text
    assert "Model calls:" not in text


def test_docx_builds_embedded_structured_math_as_semantic_omml():
    equation = thermodynamic_equation_ast()
    content = build_assessment_docx(
        run_id=7,
        prompt_id=8,
        condition_code="C111",
        run_number=1,
        course="MSE302",
        topic="Phase transformations",
        questions=[{
            "metadata": {"question_title": "Clapeyron relation"},
            "body": "Use dP/dT = Delta H / (T * Delta V) to calculate the slope.",
            "body_segments": [
                {"type": "text", "text": "Use "},
                {"type": "math", "math": equation},
                {"type": "text", "text": " to calculate the slope."},
            ],
            "options": [{
                "body": "dP/dT = Delta H / (T * Delta V)",
                "is_correct": True,
                "segments": [{"type": "math", "math": equation}],
            }],
            "model_answer": "Apply dP/dT = Delta H / (T * Delta V).",
            "model_answer_segments": [
                {"type": "text", "text": "Apply "},
                {"type": "math", "math": equation},
                {"type": "text", "text": "."},
            ],
        }],
    )

    with ZipFile(BytesIO(content)) as archive:
        document_xml = archive.read("word/document.xml")

    assert document_xml.count(b"<m:oMath>") == 3
    assert document_xml.count(b"<m:f>") == 6
    assert b"<m:num>" in document_xml
    assert b"<m:den>" in document_xml
    assert "ΔH".encode() in document_xml
    assert "ΔV".encode() in document_xml
    assert b"dP/dT" not in document_xml


def test_docx_serializes_scripts_radicals_and_matrices_to_omml():
    content = build_assessment_docx(
        run_id=9,
        prompt_id=10,
        condition_code="C010",
        run_number=1,
        course="MATH",
        topic="Structured math",
        questions=[{
            "metadata": {},
            "body": "Inspect the expressions.",
            "options": [],
            "model_answer": "See the native equations.",
            "equations": [
                {
                    "label": "Subscript",
                    "math": {
                        "type": "subscript",
                        "base": {"type": "symbol", "name": "x"},
                        "subscript": {"type": "number", "value": "1"},
                    },
                    "location": "solution",
                },
                {
                    "label": "Power",
                    "math": {
                        "type": "superscript",
                        "base": {"type": "symbol", "name": "x"},
                        "superscript": {"type": "number", "value": "2"},
                    },
                    "location": "solution",
                },
                {
                    "label": "Root",
                    "math": {
                        "type": "radical",
                        "radicand": {"type": "symbol", "name": "x"},
                    },
                    "location": "solution",
                },
                {
                    "label": "Matrix",
                    "math": {
                        "type": "matrix",
                        "rows": [
                            [
                                {"type": "number", "value": "1"},
                                {"type": "number", "value": "0"},
                            ],
                            [
                                {"type": "number", "value": "0"},
                                {"type": "number", "value": "1"},
                            ],
                        ],
                    },
                    "location": "solution",
                },
            ],
        }],
    )

    with ZipFile(BytesIO(content)) as archive:
        document_xml = archive.read("word/document.xml")

    for tag in (b"<m:sSub>", b"<m:sSup>", b"<m:rad>", b"<m:m>", b"<m:mr>"):
        assert tag in document_xml


def test_docx_preserves_fixed_five_section_template_and_solution_layout():
    content = _build_assessment_docx(
        run_id=40,
        prompt_id=41,
        condition_code="C100",
        run_number=1,
        course="MSE302",
        topic="Phase stability",
        questions=[
            {
                "type": "mcq",
                "metadata": {"question_title": "Driving force"},
                "body": "Which expression gives the driving force?",
                "options": [
                    {"body": "The Gibbs-energy change", "is_correct": True},
                    {"body": "The temperature alone", "is_correct": False},
                ],
                "model_answer": (
                    "Step 1: Write the governing relation\n"
                    "Begin with the Gibbs-energy balance.\n"
                    "[[EQ:driving_force]]\n"
                    "Therefore, the first choice is correct. Physically, the sign determines spontaneity."
                ),
                "equations": [{
                    "label": "driving_force",
                    "expression": "DeltaG = DeltaH - T DeltaS",
                    "location": "solution",
                }],
                "quality_checks": [
                    {"criterion": "Correctness", "rating": 5, "comment": "Correct relation."},
                    {"criterion": "Clarity", "rating": 4, "comment": "Clear wording."},
                ],
                "revision_options": ["Change the thermodynamic conditions."],
            },
            {
                "type": "short_answer",
                "metadata": {"question_title": "Interpretation"},
                "body": "Interpret a negative driving force.",
                "options": [],
                "model_answer": "Begin with the sign convention. Therefore, the process is spontaneous.",
                "equations": [],
                "quality_checks": [
                    {"criterion": "Correctness", "rating": 5, "comment": "Correct interpretation."},
                ],
                "revision_options": ["Ask for a physical example."],
            },
        ],
    )

    document = Document(BytesIO(content))
    paragraphs = document.paragraphs
    text = [paragraph.text for paragraph in paragraphs]
    required = [
        "1. Assessment Metadata",
        "2. Student-Facing Questions",
        "3. Fully Worked Solution",
        "4. Assessment Quality Check",
        "5. Suggested Revision Options",
    ]
    assert [text.index(heading) for heading in required] == sorted(
        text.index(heading) for heading in required
    )
    assert text.index("Question 1 — Driving force") < text.index("Question 2 — Interpretation")
    assert text.index("Question 2 — Interpretation") < text.index("Solution 1 — Driving force")
    assert text.index("Solution 1 — Driving force") < text.index("Solution 2 — Interpretation")
    assert text.index("Answer Key") < text.index("Solution 1 — Driving force")
    assert "Questions" not in text and "Solutions" not in text
    assert not any(value.lower().startswith("step 1") for value in text)
    assert len([p for p in paragraphs if p.style.name == "Solution Equation"]) == 1

    assert document.tables[0].cell(0, 0).text == "Field"
    assert document.tables[0].cell(0, 1).text == "Entry"
    metadata_rows = [tuple(cell.text for cell in row.cells) for row in document.tables[0].rows]
    assert ("MSE202 Concept(s)", "Not provided") in metadata_rows
    quality = next(table for table in document.tables if table.cell(0, 0).text == "Criterion")
    assert len(quality.rows) == 4
    assert [cell.text for cell in quality.rows[0].cells] == [
        "Criterion", "Rating / 5", "Comment", "User Rating", "User Comment"
    ]

    with ZipFile(BytesIO(content)) as archive:
        document_xml = archive.read("word/document.xml")
    assert b"[[EQ:" not in document_xml
    assert b"<m:oMath" in document_xml


@pytest.mark.parametrize(
    ("mutation", "message"),
    [
        (lambda questions: questions.clear(), "questions are required"),
        (lambda questions: questions[0].pop("metadata"), "metadata is required"),
        (lambda questions: questions[0].pop("model_answer"), "model answer is required"),
        (lambda questions: questions[0].pop("quality_checks"), "quality-check rows are required"),
        (lambda questions: questions[0].pop("revision_options"), "revision options are required"),
        (
            lambda questions: questions[0].update(model_answer="Use [[EQ:missing]]."),
            "unresolved equation reference missing",
        ),
    ],
)
def test_docx_preflight_rejects_missing_required_content(mutation, message):
    questions = [{
        "metadata": {"question_title": "Complete fixture"},
        "body": "Explain the result.",
        "options": [],
        "model_answer": "Begin with the definition. Therefore, the result follows.",
        "equations": [],
        "quality_checks": [{"criterion": "Correctness", "rating": 5, "comment": "Correct."}],
        "revision_options": ["Change the input values."],
    }]
    mutation(questions)
    with pytest.raises(DocxExportValidationError, match=message):
        _build_assessment_docx(
            run_id=50,
            prompt_id=51,
            condition_code="C010",
            run_number=1,
            course="MSE202",
            topic="Validation",
            questions=questions,
        )
