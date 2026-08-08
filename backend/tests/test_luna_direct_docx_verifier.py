from copy import deepcopy
from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from docx.oxml import OxmlElement

from backend.services.luna_direct_docx_verifier import (
    LunaDirectDocxVerifier,
    _referenced_equations,
)


ASSESSMENT = {
    "assessment_metadata": {"course": "MSE302 Thermodynamics II"},
    "questions": [
        {
            "body": "Calculate the equilibrium force.",
            "model_answer": "The equilibrium force is ten newtons.",
            "equations": [
                {"label": "force", "expression": "F = 10 N", "location": "solution"}
            ],
        }
    ]
}


def fixture_docx(
    *,
    assessment=ASSESSMENT,
    question=True,
    solution=True,
    equation=True,
    placeholder=False,
    valid_section_order=True,
    raw_linear_math=False,
    display_equation=True,
    include_metadata=True,
    repeat_metadata_header=True,
    style_metadata_first_row=True,
    include_header_footer=True,
    equation_justification="center",
    embed_placeholder_equation=False,
    image=False,
    image_alt_text="Illustration of the canonical solution relationship",
):
    document = Document()
    equation_inserted = False

    def append_equation(*, as_display):
        nonlocal equation_inserted
        container = OxmlElement("m:oMathPara") if as_display else None
        if container is not None and equation_justification is not None:
            properties = OxmlElement("m:oMathParaPr")
            justification = OxmlElement("m:jc")
            justification.set(
                "{http://schemas.openxmlformats.org/officeDocument/2006/math}val",
                equation_justification,
            )
            properties.append(justification)
            container.append(properties)
        math = OxmlElement("m:oMath")
        run = OxmlElement("m:r")
        text = OxmlElement("m:t")
        text.text = "F_1=10N" if raw_linear_math else "F=10N"
        run.append(text)
        math.append(run)
        if container is None:
            document.paragraphs[-1]._p.append(math)
        else:
            container.append(math)
            document._body._element.append(container)
        equation_inserted = True

    if include_header_footer:
        document.sections[0].header.paragraphs[0].text = "Blueprint Lab"
        footer = document.sections[0].footer.paragraphs[0]
        footer.add_run("Page ")
        for instruction in ("PAGE", "NUMPAGES"):
            field = OxmlElement("w:fldSimple")
            field.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instr", instruction)
            footer._p.append(field)
            if instruction == "PAGE":
                footer.add_run(" of ")
    if include_metadata:
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Field"
        table.rows[0].cells[1].text = "Entry"
        if repeat_metadata_header:
            row_properties = table.rows[0]._tr.get_or_add_trPr()
            row_properties.append(OxmlElement("w:tblHeader"))
        if not style_metadata_first_row:
            table_look = table._tbl.tblPr.first_child_found_in("w:tblLook")
            table_look.set(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}firstRow",
                "0",
            )
        cells = table.add_row().cells
        cells[0].text = "Course"
        cells[1].text = "MSE302 Thermodynamics II"
    if valid_section_order:
        document.add_heading("2. Student-Facing Questions", level=1)
        document.add_heading("Question 1 - Equilibrium force", level=2)
        if question:
            document.add_paragraph(assessment["questions"][0]["body"])
        document.add_heading("3. Fully Worked Solutions", level=1)
        document.add_heading("Solution 1 - Equilibrium force", level=2)
        if solution:
            answer = assessment["questions"][0]["model_answer"]
            if embed_placeholder_equation and "[[EQ:force]]" in answer:
                before, after = answer.split("[[EQ:force]]", 1)
                paragraph = document.add_paragraph()
                paragraph.add_run(before)
                if display_equation:
                    append_equation(as_display=True)
                    document.add_paragraph(after)
                else:
                    append_equation(as_display=False)
                    paragraph.add_run(after)
            else:
                document.add_paragraph(answer)
    else:
        document.add_heading("3. Fully Worked Solutions", level=1)
        document.add_heading("Solution 1 - Equilibrium force", level=2)
        if solution:
            document.add_paragraph(assessment["questions"][0]["model_answer"])
        document.add_heading("2. Student-Facing Questions", level=1)
        document.add_heading("Question 1 - Equilibrium force", level=2)
        if question:
            document.add_paragraph(assessment["questions"][0]["body"])
    if placeholder:
        document.add_paragraph("[[EQ:force]]")
    if equation and not equation_inserted:
        if not display_equation:
            document.add_paragraph()
        append_equation(as_display=display_equation)
    if image:
        paragraph = document.add_paragraph()
        drawing = OxmlElement("w:drawing")
        inline = OxmlElement("wp:inline")
        properties = OxmlElement("wp:docPr")
        properties.set("id", "1")
        properties.set("name", "Solution figure")
        if image_alt_text is not None:
            properties.set("descr", image_alt_text)
        inline.append(properties)
        drawing.append(inline)
        paragraph._p.append(drawing)
    target = BytesIO()
    document.save(target)
    return target.getvalue()


def without_main_document(content):
    source = ZipFile(BytesIO(content))
    target = BytesIO()
    with source, ZipFile(target, "w", ZIP_DEFLATED) as output:
        for item in source.infolist():
            if item.filename != "word/document.xml":
                output.writestr(item, source.read(item.filename))
    return target.getvalue()


def codes(report):
    return {issue.code for issue in report.issues}


def test_orders_equations_by_placeholder_occurrence_not_manifest_order():
    questions = [
        {
            "body": "Use [[EQ:second]] before [[EQ:first]].",
            "model_answer": "Done.",
            "equations": [
                {"label": "first", "location": "question"},
                {"label": "second", "location": "question"},
            ],
        }
    ]

    ordered = _referenced_equations(questions)

    assert [item["label"] for item in ordered] == ["second", "first"]


def test_accepts_valid_structural_canonical_and_omml_docx():
    report = LunaDirectDocxVerifier().verify(fixture_docx(), ASSESSMENT)

    assert report.valid is True
    assert report.issues == ()


def test_rejects_corrupt_zip():
    report = LunaDirectDocxVerifier().verify(b"not-a-docx", ASSESSMENT)

    assert report.valid is False
    assert codes(report) == {"docx_package_invalid"}


def test_rejects_missing_main_document():
    report = LunaDirectDocxVerifier().verify(
        without_main_document(fixture_docx()), ASSESSMENT
    )

    assert "docx_package_invalid" in codes(report)


def test_rejects_missing_canonical_question_text():
    report = LunaDirectDocxVerifier().verify(fixture_docx(question=False), ASSESSMENT)

    assert "canonical_question_missing" in codes(report)


def test_rejects_missing_canonical_solution_text():
    report = LunaDirectDocxVerifier().verify(fixture_docx(solution=False), ASSESSMENT)

    assert "canonical_solution_missing" in codes(report)


def test_rejects_unresolved_equation_placeholder():
    report = LunaDirectDocxVerifier().verify(fixture_docx(placeholder=True), ASSESSMENT)

    assert "equation_placeholder_unresolved" in codes(report)


def test_rejects_assessment_equations_without_native_omml():
    report = LunaDirectDocxVerifier().verify(fixture_docx(equation=False), ASSESSMENT)

    assert "native_equation_missing" in codes(report)


def test_rejects_interleaved_or_reversed_question_solution_sections():
    report = LunaDirectDocxVerifier().verify(
        fixture_docx(valid_section_order=False), ASSESSMENT
    )

    assert "assessment_section_order_invalid" in codes(report)


def test_rejects_raw_linear_subscript_markers_inside_native_math():
    report = LunaDirectDocxVerifier().verify(
        fixture_docx(raw_linear_math=True), ASSESSMENT
    )

    assert "native_equation_structure_invalid" in codes(report)


def test_structured_math_not_legacy_slashes_drives_fraction_requirement():
    assessment = deepcopy(ASSESSMENT)
    assessment["questions"][0]["equations"][0].update(
        {
            "expression": "J/mol",
            "math": {
                "type": "sequence",
                "items": [
                    {"type": "symbol", "name": "J"},
                    {"type": "operator", "value": "/"},
                    {"type": "symbol", "name": "mol"},
                ],
            },
        }
    )

    report = LunaDirectDocxVerifier().verify(fixture_docx(), assessment)

    assert "native_equation_structure_invalid" not in codes(report)


def test_rejects_missing_structures_required_by_math_ast_without_expression():
    assessment = deepcopy(ASSESSMENT)
    assessment["questions"][0]["equations"][0].update(
        {
            "expression": None,
            "math": {
                "type": "fraction",
                "numerator": {"type": "number", "value": "1"},
                "denominator": {"type": "symbol", "name": "x"},
            },
        }
    )

    report = LunaDirectDocxVerifier().verify(fixture_docx(), assessment)

    issue = next(
        item
        for item in report.issues
        if item.code == "native_equation_structure_invalid"
    )
    assert "fraction expected 1, found 0" in issue.evidence


def test_rejects_missing_function_and_delimiter_required_by_math_ast():
    assessment = deepcopy(ASSESSMENT)
    assessment["questions"][0]["equations"][0].update(
        {
            "expression": None,
            "math": {
                "type": "function",
                "name": "ln",
                "argument": {
                    "type": "delimiter",
                    "opening": "(",
                    "closing": ")",
                    "content": {"type": "symbol", "name": "x"},
                },
            },
        }
    )

    report = LunaDirectDocxVerifier().verify(fixture_docx(), assessment)

    issue = next(
        item
        for item in report.issues
        if item.code == "native_equation_structure_invalid"
    )
    assert "function expected 1, found 0" in issue.evidence
    assert "delimiter expected 1, found 0" in issue.evidence


def test_rejects_missing_named_greek_glyph_required_by_math_ast():
    assessment = deepcopy(ASSESSMENT)
    assessment["questions"][0]["equations"][0].update(
        {
            "expression": None,
            "math": {"type": "symbol", "name": "alpha"},
        }
    )

    report = LunaDirectDocxVerifier().verify(fixture_docx(), assessment)

    issue = next(
        item
        for item in report.issues
        if item.code == "native_equation_structure_invalid"
    )
    assert "greek glyph α" in issue.evidence


def test_rejects_legacy_display_math_without_display_container():
    report = LunaDirectDocxVerifier().verify(
        fixture_docx(display_equation=False), ASSESSMENT
    )

    assert "native_equation_placement_invalid" in codes(report)


def test_accepts_inline_math_when_placeholder_is_embedded_in_prose():
    assessment = deepcopy(ASSESSMENT)
    assessment["questions"][0]["model_answer"] = (
        "The equilibrium force is [[EQ:force]] for this case."
    )

    report = LunaDirectDocxVerifier().verify(
        fixture_docx(
            assessment=assessment,
            display_equation=False,
            embed_placeholder_equation=True,
        ),
        assessment,
    )

    assert report.valid is True


def test_rejects_embedded_inline_placeholder_promoted_to_display_math():
    assessment = deepcopy(ASSESSMENT)
    assessment["questions"][0]["model_answer"] = (
        "The equilibrium force is [[EQ:force]] for this case."
    )

    report = LunaDirectDocxVerifier().verify(
        fixture_docx(
            assessment=assessment,
            display_equation=True,
            embed_placeholder_equation=True,
        ),
        assessment,
    )

    assert "native_equation_placement_invalid" in codes(report)


def test_rejects_standalone_placeholder_demoted_to_inline_math():
    assessment = deepcopy(ASSESSMENT)
    assessment["questions"][0]["model_answer"] = (
        "The governing relation is\n[[EQ:force]]\nTherefore, the result follows."
    )

    report = LunaDirectDocxVerifier().verify(
        fixture_docx(
            assessment=assessment,
            display_equation=False,
            embed_placeholder_equation=True,
        ),
        assessment,
    )

    assert "native_equation_placement_invalid" in codes(report)


def test_accepts_omitted_math_justification_as_center_group_default():
    report = LunaDirectDocxVerifier().verify(
        fixture_docx(equation_justification=None), ASSESSMENT
    )

    assert report.valid is True


def test_accepts_solution_image_with_descriptive_alt_text():
    report = LunaDirectDocxVerifier().verify(
        fixture_docx(image=True), ASSESSMENT
    )

    assert "image_alt_text_missing" not in codes(report)


def test_rejects_solution_image_without_alt_text():
    report = LunaDirectDocxVerifier().verify(
        fixture_docx(image=True, image_alt_text=None), ASSESSMENT
    )

    assert "image_alt_text_missing" in codes(report)


def test_rejects_explicitly_left_justified_display_math():
    report = LunaDirectDocxVerifier().verify(
        fixture_docx(equation_justification="left"), ASSESSMENT
    )

    issue = next(
        item for item in report.issues if item.code == "native_equation_display_invalid"
    )
    assert "explicit justification 'left'" in issue.evidence


def test_rejects_missing_assessment_metadata_table():
    report = LunaDirectDocxVerifier().verify(
        fixture_docx(include_metadata=False), ASSESSMENT
    )

    assert "assessment_metadata_table_missing" in codes(report)


def test_accepts_metadata_first_row_styling_without_repeat_header_property():
    report = LunaDirectDocxVerifier().verify(
        fixture_docx(repeat_metadata_header=False), ASSESSMENT
    )

    assert report.valid is True


def test_rejects_metadata_table_without_any_header_row_marker():
    report = LunaDirectDocxVerifier().verify(
        fixture_docx(
            repeat_metadata_header=False,
            style_metadata_first_row=False,
        ),
        ASSESSMENT,
    )

    issue = next(
        item
        for item in report.issues
        if item.code == "assessment_metadata_header_invalid"
    )
    assert "w:tblHeader" in issue.evidence


def test_accepts_comma_separated_metadata_lists():
    assessment = deepcopy(ASSESSMENT)
    assessment["assessment_metadata"] = {
        "mse202_concepts": ["Phase equilibrium", "Chemical potential"]
    }
    document = Document(BytesIO(fixture_docx()))
    table = document.tables[0]
    table.rows[1].cells[0].text = "MSE202 Concept(s)"
    table.rows[1].cells[1].text = "Phase equilibrium, Chemical potential"
    target = BytesIO()
    document.save(target)

    report = LunaDirectDocxVerifier().verify(target.getvalue(), assessment)

    assert report.valid is True


def test_metadata_mismatch_reports_expected_label_value_and_actual_value():
    assessment = deepcopy(ASSESSMENT)
    assessment["assessment_metadata"]["course"] = "Expected Course"

    report = LunaDirectDocxVerifier().verify(fixture_docx(), assessment)

    issue = next(
        item for item in report.issues if item.code == "assessment_metadata_invalid"
    )
    assert "expected label 'Course' with value 'Expected Course'" in issue.evidence
    assert "found 'MSE302 Thermodynamics II'" in issue.evidence


def test_rejects_missing_header_and_footer():
    report = LunaDirectDocxVerifier().verify(
        fixture_docx(include_header_footer=False), ASSESSMENT
    )

    assert {"header_missing", "footer_missing"} <= codes(report)


@pytest.mark.parametrize(
    "expression",
    ["x_A", "x^2", "a/b", "ln(x)", "x_alpha"],
)
def test_rejects_math_wrappers_missing_expression_required_structure(expression):
    assessment = deepcopy(ASSESSMENT)
    assessment["questions"][0]["equations"][0]["expression"] = expression

    report = LunaDirectDocxVerifier().verify(fixture_docx(), assessment)

    assert "native_equation_structure_invalid" in codes(report)
