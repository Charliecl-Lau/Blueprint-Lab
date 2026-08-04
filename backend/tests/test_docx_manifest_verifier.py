from io import BytesIO
from types import SimpleNamespace

from docx import Document

from backend.services.docx_manifest_verifier import REQUIRED_HEADINGS, DocxManifestVerifier
from backend.tests.test_docx_authoring_schema import manifest


def rendered_manifest_docx(value):
    document = Document()
    for heading in REQUIRED_HEADINGS:
        document.add_heading(heading, level=1)
    document.add_table(rows=1, cols=2)
    document.add_table(rows=1, cols=2)
    document.add_table(rows=1, cols=5)
    for label in ("Step 1", "Final Answer", "Distractor Analysis", "Conclusion"):
        document.add_paragraph(label)
    question = value["questions"][0]
    strings = [question["title"], question["body"]]
    strings.extend(option["body"] for option in question["options"])
    solution = question["solution"]
    strings.extend([solution["governing_concept"], *solution["application_steps"], solution["conclusion"]])
    strings.extend(item["explanation"] for item in solution["option_elimination"])
    strings.extend([value["answer_key"][0]["correct_option_id"], value["answer_key"][0]["answer"]])
    strings.extend(value["revision_options"])
    for item in strings: document.add_paragraph(item)
    output = BytesIO(); document.save(output); return output.getvalue()


def test_manifest_verifier_checks_visible_content_and_source_mapping():
    value = manifest()
    grounding = SimpleNamespace(original_assessment={"questions": [{"id": "11"}]})
    report = DocxManifestVerifier().verify(rendered_manifest_docx(value), value, grounding)
    assert report.valid


def test_manifest_verifier_rejects_missing_visible_solution_content():
    value = manifest()
    grounding = SimpleNamespace(original_assessment={"questions": [{"id": "11"}]})
    document = Document(); document.add_paragraph("Assessment Metadata")
    output = BytesIO(); document.save(output)
    assert not DocxManifestVerifier().verify(output.getvalue(), value, grounding).valid
