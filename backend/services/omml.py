import re
from typing import Iterable, Optional

from docx.oxml import OxmlElement


_GREEK_NAMES = {
    "Alpha": "Α",
    "Beta": "Β",
    "Gamma": "Γ",
    "Delta": "Δ",
    "Theta": "Θ",
    "Lambda": "Λ",
    "Mu": "Μ",
    "Pi": "Π",
    "Sigma": "Σ",
    "Phi": "Φ",
    "Psi": "Ψ",
    "Omega": "Ω",
    "alpha": "α",
    "beta": "β",
    "gamma": "γ",
    "delta": "δ",
    "epsilon": "ε",
    "theta": "θ",
    "lambda": "λ",
    "mu": "μ",
    "nu": "ν",
    "pi": "π",
    "rho": "ρ",
    "sigma": "σ",
    "tau": "τ",
    "phi": "φ",
    "chi": "χ",
    "psi": "ψ",
    "omega": "ω",
}

_LINEAR_TOKEN_PATTERN = re.compile(
    r"sqrt|[A-Za-z]+|[0-9]+(?:\.[0-9]+)?|[^\s]",
    re.IGNORECASE,
)
_EQUATION_PLACEHOLDER_PATTERN = re.compile(r"\[\[EQ:([A-Za-z0-9_-]+)\]\]")


def _sequence(items: list[dict]) -> dict:
    if len(items) == 1:
        return items[0]
    return {"type": "sequence", "items": items}


class _LinearEquationParser:
    def __init__(self, expression: str):
        self.tokens = _LINEAR_TOKEN_PATTERN.findall(expression)
        self.position = 0

    def parse(self) -> dict:
        if not self.tokens:
            raise ValueError("equation expression is empty")
        node = self._parse_equation()
        if self.position != len(self.tokens):
            raise ValueError(f"unexpected token: {self._peek()}")
        return node

    def _peek(self):
        return self.tokens[self.position] if self.position < len(self.tokens) else None

    def _take(self) -> str:
        token = self._peek()
        if token is None:
            raise ValueError("unexpected end of expression")
        self.position += 1
        return token

    def _parse_equation(self) -> dict:
        items = [self._parse_sum()]
        while self._peek() == "=":
            self._take()
            items.extend([
                {"type": "operator", "value": "="},
                self._parse_sum(),
            ])
        return _sequence(items)

    def _parse_sum(self) -> dict:
        items = [self._parse_fraction()]
        while self._peek() in ("+", "-"):
            operator = self._take()
            items.extend([
                {"type": "operator", "value": operator},
                self._parse_fraction(),
            ])
        return _sequence(items)

    def _parse_fraction(self) -> dict:
        node = self._parse_product()
        while self._peek() == "/":
            self._take()
            node = {
                "type": "fraction",
                "numerator": node,
                "denominator": self._parse_product(),
            }
        return node

    def _parse_product(self) -> dict:
        items = [self._parse_script()]
        while True:
            token = self._peek()
            if token is None or token in ("+", "-", "=", "/", ")", "}"):
                break
            if token in ("*", "·", "×"):
                items.append({"type": "operator", "value": self._take()})
                items.append(self._parse_script())
                continue
            items.append(self._parse_script())
        return _sequence(items)

    def _parse_script(self) -> dict:
        node = self._parse_atom()
        while self._peek() in ("_", "^"):
            operator = self._take()
            script = self._parse_script_value()
            if operator == "_":
                node = {"type": "subscript", "base": node, "subscript": script}
            else:
                node = {"type": "superscript", "base": node, "superscript": script}
        return node

    def _parse_script_value(self) -> dict:
        sign = self._peek()
        if sign not in ("+", "-", "−"):
            return self._parse_atom()
        self._take()
        return _sequence([
            {"type": "operator", "value": sign},
            self._parse_atom(),
        ])

    def _parse_atom(self) -> dict:
        token = self._take()
        if token in ("(", "{"):
            closing = ")" if token == "(" else "}"
            node = self._parse_equation()
            if self._take() != closing:
                raise ValueError(f"missing closing {closing}")
            return node
        if token.lower() == "sqrt" or token == "√":
            return {"type": "radical", "radicand": self._parse_atom()}
        if token in (")", "}", "_", "^", "/"):
            raise ValueError(f"unexpected token: {token}")
        if re.fullmatch(r"[0-9]+(?:\.[0-9]+)?", token):
            return {"type": "number", "value": token}
        if token in ("+", "-", "*", "·", "×", "=", ","):
            return {"type": "operator", "value": token}
        return {"type": "symbol", "name": token}


def parse_linear_expression(expression: str) -> dict:
    """Parse the supported Word-linear subset, falling back to editable math text."""
    try:
        return _LinearEquationParser(expression).parse()
    except (TypeError, ValueError):
        return {"type": "text", "text": expression}


def _as_dict(node):
    return node.model_dump(exclude_none=True) if hasattr(node, "model_dump") else node


def _math_run(value: str):
    run = OxmlElement("m:r")
    text = OxmlElement("m:t")
    text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text.text = value
    run.append(text)
    return run


def _symbol_text(name: str) -> str:
    if name in _GREEK_NAMES:
        return _GREEK_NAMES[name]
    for word in sorted(_GREEK_NAMES, key=len, reverse=True):
        if name.startswith(word) and len(name) > len(word):
            return f"{_GREEK_NAMES[word]}{name[len(word):]}"
    return name


def _container(tag: str, children: Iterable):
    element = OxmlElement(tag)
    for child in children:
        element.append(child)
    return element


def _serialize(node) -> list:
    node = _as_dict(node)
    kind = node["type"]
    if kind == "text":
        return [_math_run(node["text"])]
    if kind == "symbol":
        return [_math_run(_symbol_text(node["name"]))]
    if kind == "number":
        return [_math_run(node["value"])]
    if kind == "operator":
        return [_math_run(node["value"])]
    if kind == "differential":
        return [_math_run(f"d{_symbol_text(node['variable'])}")]
    if kind == "sequence":
        return [child for item in node["items"] for child in _serialize(item)]
    if kind == "equation":
        return _serialize(node["left"]) + [_math_run("=")] + _serialize(node["right"])
    if kind == "product":
        operator = {"implicit": "", "dot": "·", "cross": "×"}.get(
            node.get("operator", "implicit"), ""
        )
        result = []
        for index, term in enumerate(node["terms"]):
            if index and operator:
                result.append(_math_run(operator))
            result.extend(_serialize(term))
        return result
    if kind == "fraction":
        fraction = OxmlElement("m:f")
        fraction.append(_container("m:num", _serialize(node["numerator"])))
        fraction.append(_container("m:den", _serialize(node["denominator"])))
        return [fraction]
    if kind in ("subscript", "superscript"):
        element = OxmlElement("m:sSub" if kind == "subscript" else "m:sSup")
        element.append(_container("m:e", _serialize(node["base"])))
        script_key = "subscript" if kind == "subscript" else "superscript"
        script_tag = "m:sub" if kind == "subscript" else "m:sup"
        element.append(_container(script_tag, _serialize(node[script_key])))
        return [element]
    if kind == "radical":
        radical = OxmlElement("m:rad")
        degree = node.get("degree")
        properties = OxmlElement("m:radPr")
        if degree is None:
            degree_hidden = OxmlElement("m:degHide")
            degree_hidden.set("{http://schemas.openxmlformats.org/officeDocument/2006/math}val", "1")
            properties.append(degree_hidden)
        radical.append(properties)
        radical.append(_container("m:deg", _serialize(degree) if degree else []))
        radical.append(_container("m:e", _serialize(node["radicand"])))
        return [radical]
    if kind == "matrix":
        matrix = OxmlElement("m:m")
        for row in node["rows"]:
            matrix_row = OxmlElement("m:mr")
            for cell in row:
                matrix_row.append(_container("m:e", _serialize(cell)))
            matrix.append(matrix_row)
        return [matrix]
    raise ValueError(f"Unsupported structured math node: {kind}")


def append_math(paragraph, node) -> None:
    math = OxmlElement("m:oMath")
    for child in _serialize(node):
        math.append(child)
    paragraph._p.append(math)


def append_linear_math(paragraph, expression: str) -> None:
    append_math(paragraph, parse_linear_expression(expression))


def append_text_with_equations(
    paragraph,
    text: str,
    equations: list[dict],
    location: str,
) -> set[str]:
    # An explicit placeholder determines placement. Location is only needed when
    # an unreferenced legacy equation is rendered as a standalone paragraph.
    equation_by_label = {equation["label"]: equation for equation in equations}
    rendered_labels = set()
    cursor = 0
    for match in _EQUATION_PLACEHOLDER_PATTERN.finditer(text):
        paragraph.add_run(text[cursor:match.start()])
        equation = equation_by_label.get(match.group(1))
        if equation is None:
            paragraph.add_run(match.group(0))
        else:
            if equation.get("math"):
                append_math(paragraph, equation["math"])
            else:
                append_linear_math(paragraph, equation["expression"])
            rendered_labels.add(equation["label"])
        cursor = match.end()
    paragraph.add_run(text[cursor:])
    return rendered_labels


def append_content(
    paragraph,
    segments,
    fallback: str,
    *,
    equations: Optional[list[dict]] = None,
    location: Optional[str] = None,
) -> set[str]:
    if not segments:
        if equations is not None and location is not None:
            return append_text_with_equations(paragraph, fallback, equations, location)
        paragraph.add_run(fallback)
        return set()
    for segment in segments:
        segment = _as_dict(segment)
        if segment["type"] == "text":
            paragraph.add_run(segment["text"])
        else:
            append_math(paragraph, segment["math"])
    return set()
