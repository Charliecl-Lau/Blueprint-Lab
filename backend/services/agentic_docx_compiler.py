"""Trusted compiler from replayable workspace IR to safe OOXML."""

from __future__ import annotations

import hashlib
import io
import re
from dataclasses import dataclass
from typing import Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from backend.services.docx_content_catalog import DocxContentCatalog
from backend.services.docx_exporter import validate_assessment_for_docx
from backend.services.docx_layout_manifest import MANIFEST_VERSION, manifest_sha256
from backend.services.docx_package_verifier import DocxPackageVerifier
from backend.services.docx_tool_workspace import DocxWorkspace, WorkspaceError
from backend.services.omml import append_content, append_linear_math, append_math


class AgenticDocxCompileError(ValueError):
    pass


@dataclass(frozen=True)
class CompiledAgenticDocx:
    docx_bytes: bytes
    docx_sha256: str
    assessment_json: dict
    layout_manifest: dict
    layout_manifest_sha256: str


class AgenticDocxCompiler:
    VERSION = "agentic-docx-compiler-v4-fixed-assessment-template"

    NAVY = "1F4E79"
    PALE_BLUE = "EAF2F8"
    GRAY = "666666"
    SECTION_TITLES = {
        "assessment_metadata": "1. Assessment Metadata",
        "questions": "2. Student-Facing Questions",
        "solutions": "3. Fully Worked Solution",
        "answer_key": None,
        "quality_check": "4. Assessment Quality Check",
        "revision_options": "5. Suggested Revision Options",
    }

    def __init__(self, package_verifier: Optional[DocxPackageVerifier] = None):
        self.package_verifier = package_verifier or DocxPackageVerifier()

    def compile(self, workspace: DocxWorkspace, *, session_id: Optional[int] = None, iteration_number: int = 0) -> CompiledAgenticDocx:
        workspace.validate_structure(require_complete=True)
        state = workspace.to_dict()
        assessment = workspace.catalog.clone_assessment()
        validate_assessment_for_docx(assessment.get("questions") or [])
        document = Document()
        self._configure(document, state, workspace.catalog)
        placements = []
        mappings = []
        rendered_equations = set()
        catalog = workspace.catalog

        # The application owns the assessment topology. Model-authored workspace
        # blocks may influence approved design tokens, but never section order,
        # headings, or assessed-content placement.
        for role in (
            "assessment_metadata",
            "questions",
            "solutions",
            "quality_check",
            "revision_options",
        ):
            self._render_block(
                document,
                catalog,
                {
                    "id": f"canonical-section-{role}",
                    "type": "section",
                    "role": role,
                },
                mappings,
                placements,
                rendered_equations,
            )
            if role == "questions":
                for index, question_id in enumerate(catalog.question_ids):
                    self._render_block(
                        document,
                        catalog,
                        {
                            "id": f"canonical-question-{question_id}",
                            "type": "question",
                            "question_id": question_id,
                            "_first_question": index == 0,
                        },
                        mappings,
                        placements,
                        rendered_equations,
                    )
            elif role == "solutions":
                if any(
                    catalog.resolve_question(qid).get("options")
                    for qid in catalog.question_ids
                ):
                    self._render_block(
                        document,
                        catalog,
                        {"id": "canonical-answer-key", "type": "answer_key"},
                        mappings,
                        placements,
                        rendered_equations,
                    )
                for question_id in catalog.question_ids:
                    self._render_block(
                        document,
                        catalog,
                        {
                            "id": f"canonical-solution-{question_id}",
                            "type": "solution",
                            "question_id": question_id,
                        },
                        mappings,
                        placements,
                        rendered_equations,
                    )
            elif role == "quality_check":
                self._render_block(
                    document,
                    catalog,
                    {"id": "canonical-quality-check", "type": "quality_check"},
                    mappings,
                    placements,
                    rendered_equations,
                )
            elif role == "revision_options":
                self._render_revision_options(document, catalog, mappings)

        missing_equations = set(workspace.catalog.equation_ids) - rendered_equations
        if missing_equations:
            raise AgenticDocxCompileError(
                f"missing equation placements: {sorted(missing_equations)}"
            )

        stream = io.BytesIO()
        document.save(stream)
        docx_bytes = stream.getvalue()
        package_report = self.package_verifier.verify(docx_bytes)
        if not package_report.valid:
            raise AgenticDocxCompileError("compiled DOCX failed package verification")
        docx_hash = hashlib.sha256(docx_bytes).hexdigest()
        manifest = {
            "schema_version": MANIFEST_VERSION,
            "compiler_version": self.VERSION,
            "workspace_hash": workspace.sha256,
            "catalog_hash": workspace.catalog.sha256,
            "tool_session_id": session_id,
            "iteration_number": iteration_number,
            "docx_hash": docx_hash,
            "block_mappings": mappings,
            "equation_placements": placements,
            "style_tokens": sorted({item.get("style") for item in state["blocks"] if item.get("style")}),
            "render_hashes": [],
            "validator_versions": package_report.tool_versions,
        }
        return CompiledAgenticDocx(
            docx_bytes, docx_hash, assessment, manifest, manifest_sha256(manifest)
        )

    def _configure(self, document, state, catalog):
        color = RGBColor(31, 78, 121)
        normal = document.styles["Normal"]
        normal.font.name = "Aptos"; normal.font.size = Pt(10)
        normal.paragraph_format.space_after = Pt(5)
        normal.paragraph_format.line_spacing = 1.08
        for index, size, before, after in ((1, 15, 8, 4), (2, 12, 6, 4), (3, 10.5, 6, 4)):
            style = document.styles[f"Heading {index}"]
            style.font.name = "Aptos Display"; style.font.size = Pt(size)
            style.font.bold = True; style.font.color.rgb = color
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True
        self._configure_solution_styles(document)
        layout = state.get("page_layout") or {}
        for section in document.sections:
            if layout.get("orientation") == "landscape":
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width, section.page_height = section.page_height, section.page_width
            margin = float(layout.get("margin_inches", 0.72))
            if not 0.4 <= margin <= 2.0:
                raise AgenticDocxCompileError("margin outside approved range")
            section.left_margin = section.right_margin = Inches(margin)
            section.top_margin = Inches(0.68)
            section.bottom_margin = Inches(0.65)
            section.header_distance = Inches(0.28)
            section.footer_distance = Inches(0.30)
            self._configure_header(section, catalog)
            self._configure_footer(section)

    @staticmethod
    def _configure_solution_styles(document):
        styles = document.styles
        for name in ("Solution Subheading", "Solution Equation", "Solution Body"):
            try:
                style = styles[name]
            except KeyError:
                style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = "Aptos"
            style.font.size = Pt(10)
            style.base_style = styles["Normal"]

        subheading = styles["Solution Subheading"]
        subheading.font.bold = True
        subheading.font.color.rgb = RGBColor(0, 0, 0)
        subheading.paragraph_format.space_before = Pt(8)
        subheading.paragraph_format.space_after = Pt(3)
        subheading.paragraph_format.keep_with_next = True

        equation = styles["Solution Equation"]
        equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        equation.paragraph_format.space_before = Pt(4)
        equation.paragraph_format.space_after = Pt(7)
        equation.paragraph_format.keep_with_next = True

        body = styles["Solution Body"]
        body.paragraph_format.space_before = Pt(0)
        body.paragraph_format.space_after = Pt(5)
        body.paragraph_format.line_spacing = 1.08

    def _configure_header(self, section, catalog):
        assessment = catalog.clone_assessment()
        metadata = assessment.get("metadata") or {}
        first = (assessment.get("questions") or [{}])[0].get("metadata") or {}
        course = metadata.get("course") or metadata.get("course_name") or first.get("course") or "Assessment"
        paragraph = section.header.paragraphs[0]
        paragraph.text = f"{course}  |  Blueprint Solutions Question Bank"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_after = Pt(3)
        for run in paragraph.runs:
            run.font.name = "Aptos"; run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(102, 102, 102)
        p_pr = paragraph._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        for key, value in (("val", "single"), ("sz", "4"), ("space", "1"), ("color", "B7B7B7")):
            bottom.set(qn(f"w:{key}"), value)
        borders.append(bottom); p_pr.append(borders)

    def _configure_footer(self, section):
        paragraph = section.footer.paragraphs[0]
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run("Page ")
        run.font.name = "Aptos"; run.font.size = Pt(8); run.font.color.rgb = RGBColor(102, 102, 102)
        self._append_field(run, "PAGE")
        run = paragraph.add_run(" of ")
        run.font.name = "Aptos"; run.font.size = Pt(8); run.font.color.rgb = RGBColor(102, 102, 102)
        self._append_field(run, "NUMPAGES")

    @staticmethod
    def _append_field(run, field):
        begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText"); instruction.set(qn("xml:space"), "preserve"); instruction.text = f" {field} "
        separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
        end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
        run._r.extend([begin, instruction, separate, end])

    def _render_block(self, document, catalog, block, mappings, placements, rendered_equations):
        kind = block["type"]
        mapping = {"block_id": block["id"], "type": kind}
        if block.get("content_ref"):
            mapping["content_ref"] = block["content_ref"]
        if block.get("question_id") is not None:
            mapping["question_id"] = str(block["question_id"])
        mappings.append(mapping)
        if kind == "section":
            role = block.get("role")
            title = self.SECTION_TITLES.get(role)
            if title:
                document.add_heading(title, level=1)
            elif role not in self.SECTION_TITLES:
                title = block.get("literal_text") or str(role or "Section").replace("_", " ").title()
                document.add_heading(title, level=1)
            if role == "assessment_metadata":
                self._render_metadata_table(document, catalog)
        elif kind == "heading":
            if block.get("_section_role") == "assessment_metadata":
                return
            text = catalog.resolve_text(block["content_ref"]) if block.get("content_ref") else block.get("literal_text", "")
            level = {"heading_1": 1, "heading_2": 2, "heading_3": 3}.get(block.get("style"), int(block.get("level", 2)))
            document.add_heading(text, level=max(1, min(3, level)))
        elif kind == "content":
            if block.get("_section_role") == "assessment_metadata":
                return
            paragraph = document.add_paragraph(catalog.resolve_text(block["content_ref"]))
            if block.get("_section_role") == "revision_options":
                paragraph.style = document.styles["List Number"]
        elif kind == "callout":
            paragraph = document.add_paragraph(block.get("literal_text", ""))
            paragraph.style = document.styles["Intense Quote"]
        elif kind == "page_break":
            document.add_page_break()
        elif kind == "question":
            self._render_question(
                document, catalog, str(block["question_id"]), placements,
                rendered_equations, first=bool(block.get("_first_question")),
            )
        elif kind == "solution":
            self._render_solution(document, catalog, str(block["question_id"]), placements, rendered_equations)
        elif kind == "equation":
            key = (str(block["question_id"]), block["equation_id"])
            if key in rendered_equations:
                raise AgenticDocxCompileError(f"duplicate equation placement: {key}")
            equation = catalog.resolve_equation(*key)
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            append_math(paragraph, equation["math"]) if equation.get("math") else append_linear_math(paragraph, equation["expression"])
            rendered_equations.add(key); placements.append({"block_id": block["id"], "question_id": key[0], "equation_id": key[1], "location": equation["location"]})
        elif kind == "answer_key":
            document.add_heading("Answer key", level=2)
            table = document.add_table(rows=1, cols=2)
            self._format_table(table, (1.30, 1.60))
            self._header(table.rows[0], ["Question", "Correct choice"])
            for qid in catalog.question_ids:
                question = catalog.resolve_question(qid); cells = table.add_row().cells
                correct = self._correct_choice(question)
                cells[0].text = str(list(catalog.question_ids).index(qid) + 1); cells[1].text = correct
                self._format_body_row(cells, alternate=len(table.rows) % 2 == 1, centered=True)
        elif kind == "quality_check":
            table = document.add_table(rows=1, cols=5)
            self._format_table(table, (1.45, 0.62, 2.95, 0.72, 1.16))
            self._header(table.rows[0], ["Criterion", "Rating\n/ 5", "Comment", "User\nRating", "User Comment"])
            for qid in catalog.question_ids:
                for check in catalog.resolve_question(qid).get("quality_checks", []):
                    cells = table.add_row().cells
                    values = [check.get("criterion", ""), check.get("rating", ""), check.get("comment", ""), "", ""]
                    for cell, value in zip(cells, values): cell.text = str(value)
                    self._format_body_row(cells, alternate=len(table.rows) % 2 == 1, small=True)
        elif kind == "table":
            refs = block.get("content_refs") or []
            columns = max(1, int(block.get("columns", 2)))
            table = document.add_table(rows=0, cols=columns); table.style = "Table Grid"
            for start in range(0, len(refs), columns):
                cells = table.add_row().cells
                for cell, ref in zip(cells, refs[start:start + columns]): cell.text = catalog.resolve_text(ref)

    def _header(self, row, values):
        row_property = row._tr.get_or_add_trPr(); repeat = OxmlElement("w:tblHeader"); repeat.set(qn("w:val"), "true"); row_property.append(repeat)
        for cell, value in zip(row.cells, values):
            cell.text = value
            self._shade(cell, self.NAVY)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0); paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.bold = True; run.font.name = "Aptos"; run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)

    def _render_metadata_table(self, document, catalog):
        assessment = catalog.clone_assessment()
        questions = assessment.get("questions") or []
        top = assessment.get("metadata") or {}

        def values(*keys):
            found = []
            for source in [top, *((question.get("metadata") or {}) for question in questions)]:
                for key in keys:
                    value = source.get(key)
                    if isinstance(value, list):
                        candidates = [str(item) for item in value if item not in (None, "")]
                    elif value not in (None, ""):
                        candidates = [str(value)]
                    else:
                        candidates = []
                    for candidate in candidates:
                        if candidate not in found:
                            found.append(candidate)
            return "; ".join(found)

        course = values("course", "course_name")
        item_count = len(questions)
        supplied_title = values("assessment_title", "question_bank_title")
        title = supplied_title or (
            f"{course} — {item_count}-Item Assessment" if course else f"{item_count}-Item Assessment"
        )
        rows = [
            ("Question title", title),
            ("Course", course),
            ("Topic", values("topic")),
            ("Question type", values("question_type")),
            ("Difficulty level", values("difficulty_level")),
            ("Intended assessment setting", values("intended_assessment_setting")),
            ("MSE202 concept(s) used", values("mse202_concepts")),
            ("MSE302 concept(s) used", values("mse302_concepts")),
            ("Concept-map bridge", values("concept_map_bridge")),
            ("Materials science context", values("materials_science_context")),
            ("Numerical computation", values("numerical_computation")),
            ("Estimated time for a well-prepared student", values("estimated_time", "estimated_time_minutes")),
            ("Learning objective(s)", values("learning_objectives")),
        ]
        rows = [
            (label, value or "Not provided")
            for label, value in rows
        ]
        table = document.add_table(rows=1, cols=2)
        self._format_table(table, (1.72, 5.20))
        self._header(table.rows[0], ["Field", "Entry"])
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = label; cells[1].text = value
            self._format_body_row(cells, alternate=len(table.rows) % 2 == 1)
            for run in cells[0].paragraphs[0].runs:
                run.bold = True

    def _format_table(self, table, widths):
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        width_values = [Inches(value) for value in widths]
        table_width = table._tbl.tblPr.first_child_found_in("w:tblW")
        if table_width is not None:
            table_width.set(qn("w:w"), str(sum(width.twips for width in width_values)))
            table_width.set(qn("w:type"), "dxa")
        for grid_column, width in zip(table._tbl.tblGrid.gridCol_lst, width_values):
            grid_column.set(qn("w:w"), str(width.twips))
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), "nil")
            borders.append(element)
        table._tbl.tblPr.append(borders)
        for row in table.rows:
            for cell, width in zip(row.cells, width_values):
                cell.width = width

    def _format_body_row(self, cells, *, alternate=False, centered=False, small=False):
        for index, cell in enumerate(cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if alternate:
                self._shade(cell, self.PALE_BLUE)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0 if small else 1.05
            if centered or index in ({1, 3} if small else set()):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = "Aptos"
                if small:
                    run.font.size = Pt(7.5)

    @staticmethod
    def _shade(cell, color):
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color)
        cell._tc.get_or_add_tcPr().append(shading)

    @staticmethod
    def _correct_choice(question):
        for index, option in enumerate(question.get("options") or []):
            if option.get("is_correct") is True:
                return chr(65 + index)
        return "See solution"

    @staticmethod
    def _render_revision_options(document, catalog, mappings):
        for question_index, question_id in enumerate(catalog.question_ids, start=1):
            question = catalog.resolve_question(question_id)
            for revision_index, revision in enumerate(
                question.get("revision_options") or [], start=1
            ):
                paragraph = document.add_paragraph(style="List Number")
                paragraph.add_run(f"Q{question_index}: ").bold = True
                paragraph.add_run(str(revision))
                mappings.append({
                    "block_id": (
                        f"canonical-revision-{question_id}-{revision_index}"
                    ),
                    "type": "content",
                    "question_id": str(question_id),
                    "content_ref": (
                        f"question.{question_id}.revision_option."
                        f"{revision_index - 1}"
                    ),
                })

    def _render_question(self, document, catalog, qid, placements, rendered, *, first=False):
        question = catalog.resolve_question(qid)
        if first and question.get("type") == "mcq":
            intro = document.add_paragraph()
            intro.paragraph_format.keep_with_next = True
            intro.add_run("Each question has ")
            intro.add_run("one best answer").bold = True
            intro.add_run(".")
        title = (question.get("metadata") or {}).get("question_title")
        number = list(catalog.question_ids).index(qid) + 1
        document.add_heading(f"Question {number}" + (f" — {title}" if title else ""), level=2)
        paragraph = document.add_paragraph()
        labels = append_content(paragraph, question.get("body_segments"), question.get("body", ""), equations=question.get("equations", []), location="question")
        for index, option in enumerate(question.get("options", [])):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.05)
            paragraph.paragraph_format.first_line_indent = Inches(0)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.add_run(f"{chr(65 + index)}. ").bold = True
            labels.update(append_content(paragraph, option.get("segments"), option.get("body", ""), equations=question.get("equations", []), location="question"))
        self._record_inline(qid, labels, question, "question", placements, rendered)
        self._render_remaining_equations(
            document, qid, question, "question", placements, rendered
        )

    def _render_solution(self, document, catalog, qid, placements, rendered):
        question = catalog.resolve_question(qid)
        title = (question.get("metadata") or {}).get("question_title")
        number = list(catalog.question_ids).index(qid) + 1
        document.add_heading(f"Solution {number}" + (f" — {title}" if title else ""), level=2)
        answer = question.get("model_answer")
        if not answer:
            answer = next((o.get("body", "") for o in question.get("options", []) if o.get("is_correct")), "")
        if question.get("type") == "mcq":
            correct = document.add_paragraph()
            correct.paragraph_format.keep_with_next = True
            correct.add_run(f"Correct answer: {self._correct_choice(question)}.").bold = True
        labels = self._render_guided_solution(document, question, answer)
        self._record_inline(qid, labels, question, "solution", placements, rendered)
        self._render_remaining_equations(
            document, qid, question, "solution", placements, rendered
        )

    def _render_guided_solution(self, document, question, answer):
        labels = set()
        lines = answer.splitlines()
        has_step_labels = any(
            re.match(r"^\s*(?:\*\*)?Step\s+\d+\s+[—-]", line)
            for line in lines
        )
        if len(lines) == 1 and not has_step_labels:
            lines = [
                part.strip()
                for part in re.split(r"(?<=[.!?])\s+(?=[A-Z])", answer)
                if part.strip()
            ]

        for source_line in lines or [answer]:
            line = source_line.strip()
            if not line:
                spacer = document.add_paragraph()
                spacer.paragraph_format.space_after = Pt(0)
                continue
            if line.startswith("**") and line.endswith("**") and len(line) > 4:
                line = line[2:-2].strip()
            if question.get("type") == "mcq" and re.match(
                r"^Correct answer\s*:", line, flags=re.IGNORECASE
            ):
                continue

            equation_only = re.fullmatch(
                r"\[\[EQ:[A-Za-z0-9_-]+\]\][.!]?", line
            )
            is_step = bool(re.match(r"^Step\s+\d+\s+[—-]\s+", line))
            is_distractor_heading = line.rstrip(":").casefold() == (
                "why the other choices are incorrect"
            )
            distractor = re.match(r"^([A-E])\.\s+(.*)$", line)

            if equation_only:
                paragraph = document.add_paragraph(style="Solution Equation")
                labels.update(append_content(
                    paragraph, None, line.rstrip("."),
                    equations=question.get("equations", []), location="solution",
                ))
                continue
            if is_step:
                # Step labels are decorative scaffolding. Suppress them so legacy
                # responses still read as a continuous guided derivation.
                continue
            if is_distractor_heading:
                paragraph = document.add_paragraph(style="Solution Subheading")
                paragraph.add_run(line).bold = True
                continue
            paragraph = document.add_paragraph(style="Solution Body")
            if distractor:
                paragraph.add_run(f"{distractor.group(1)}. ").bold = True
                line = distractor.group(2)
            labels.update(append_content(
                paragraph, None, line,
                equations=question.get("equations", []), location="solution",
            ))
        return labels

    @staticmethod
    def _render_remaining_equations(document, qid, question, location, placements, rendered):
        for equation in question.get("equations", []):
            equation_id = equation.get("equation_id") or equation.get("label")
            key = (qid, equation_id)
            if equation.get("location") != location or key in rendered:
                continue
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if equation.get("math"):
                append_math(paragraph, equation["math"])
            else:
                append_linear_math(paragraph, equation["expression"])
            placements.append({
                "block_id": None,
                "question_id": qid,
                "equation_id": equation_id,
                "location": location,
            })
            rendered.add(key)

    @staticmethod
    def _record_inline(qid, labels, question, location, placements, rendered):
        for label in labels:
            key = (qid, label)
            if key in rendered:
                raise AgenticDocxCompileError(f"duplicate equation placement: {key}")
            equation = next(item for item in question.get("equations", []) if item.get("label") == label)
            placements.append({"block_id": None, "question_id": qid, "equation_id": label, "location": location})
            rendered.add(key)
