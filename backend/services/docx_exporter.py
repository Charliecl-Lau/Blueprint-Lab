from io import BytesIO
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from backend.services.omml import append_content, append_linear_math, append_math


METADATA_FIELDS = (
    ("Question Title", "question_title"),
    ("Question Type", "question_type"),
    ("Difficulty Level", "difficulty_level"),
    ("MSE202 Concept(s)", "mse202_concepts"),
    ("MSE302 Concept(s)", "mse302_concepts"),
    ("Concept-Map Bridge", "concept_map_bridge"),
    ("Materials Science Context", "materials_science_context"),
    ("Estimated Time (minutes)", "estimated_time_minutes"),
    ("Learning Objectives", "learning_objectives"),
)

_EQUATION_REFERENCE = re.compile(r"\[\[EQ:([A-Za-z0-9_-]+)\]\]")


class DocxExportValidationError(ValueError):
    """Raised when required assessment content cannot fill the DOCX template."""


def _configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, before, after in (
        ("Heading 1", 18, 18, 8),
        ("Heading 2", 14, 12, 6),
        ("Heading 3", 12, 13, 5),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    choices = styles["List Bullet"]
    choices.font.name = "Arial"
    choices.font.size = Pt(11)
    choices.paragraph_format.left_indent = Inches(0.3)
    choices.paragraph_format.first_line_indent = Inches(-0.18)
    choices.paragraph_format.space_after = Pt(5)

    for name in ("Solution Body", "Solution Equation"):
        try:
            style = styles[name]
        except KeyError:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Arial"
        style.font.size = Pt(11)
        style.base_style = normal
    solution_body = styles["Solution Body"]
    solution_body.paragraph_format.space_before = Pt(0)
    solution_body.paragraph_format.space_after = Pt(6)
    solution_body.paragraph_format.line_spacing = 1.08
    solution_equation = styles["Solution Equation"]
    solution_equation.paragraph_format.space_before = Pt(3)
    solution_equation.paragraph_format.space_after = Pt(6)
    solution_equation.paragraph_format.line_spacing = 1.0
    solution_equation.paragraph_format.keep_with_next = True


def _text_value(value):
    if isinstance(value, list):
        return ", ".join(str(item) for item in value if item)
    return str(value) if value not in (None, "") else ""


def _metadata_rows(*, run_id, prompt_id, condition_code, run_number,
                   course, topic, questions, traceability=None):
    traceability = traceability or {}
    rows = [
        ("Experiment ID", _text_value(traceability.get("experiment_id")) or "Not Assigned"),
        ("Condition ID", _text_value(traceability.get("condition_id")) or "Not Assigned"),
        ("Run ID", str(run_id)),
        ("Prompt ID", str(prompt_id)),
        ("Assessment ID", _text_value(traceability.get("assessment_id")) or "Not Assigned"),
        ("Condition Code", str(condition_code)),
        ("Run Number", str(run_number)),
        ("Course", str(course)),
        ("Topic", str(topic)),
    ]
    metadata = questions[0].get("metadata", {})
    for label, key in METADATA_FIELDS:
        value = _text_value(metadata.get(key))
        rows.append((label, value or "Not provided"))
    return rows


def _set_cell_margins(cell, *, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start),
                        ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_width = table._tbl.tblPr.first_child_found_in("w:tblW")
    table_width.set(qn("w:w"), str(sum(width.twips for width in widths)))
    table_width.set(qn("w:type"), "dxa")
    for grid_column, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid_column.set(qn("w:w"), str(width.twips))


def _format_metadata_cell(cell, *, width, is_label):
    cell.width = width
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    if is_label:
        paragraph.runs[0].bold = True
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "D9EAF7")
        cell._tc.get_or_add_tcPr().append(shading)


def _add_metadata_table(document, **metadata_inputs):
    rows = _metadata_rows(**metadata_inputs)
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    widths = (Inches(2.05), Inches(4.45))
    _set_table_geometry(table, widths)

    header = table.add_row().cells
    header[0].text = "Field"
    header[1].text = "Entry"
    _format_metadata_cell(header[0], width=widths[0], is_label=True)
    _format_metadata_cell(header[1], width=widths[1], is_label=True)

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        _format_metadata_cell(cells[0], width=widths[0], is_label=True)
        _format_metadata_cell(cells[1], width=widths[1], is_label=False)


def _add_item_heading(document, *, kind, index, question):
    title = question.get("metadata", {}).get("question_title")
    text = f"{kind} {index}"
    if title:
        text += f" — {title}"
    question_id = question.get("traceability", {}).get(
        "assessment_question_id"
    )
    if question_id is not None:
        text += f" [Question ID {question_id}]"
    document.add_heading(text, level=2)


def _add_standalone_equation(document, equation):
    paragraph = document.add_paragraph(style="Solution Equation")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if equation.get("math"):
        append_math(paragraph, equation["math"])
    else:
        append_linear_math(paragraph, equation["expression"])


def _question_title(question, index):
    return question.get("metadata", {}).get("question_title") or f"Question {index}"


def _correct_choice(question):
    for index, option in enumerate(question.get("options", [])):
        if option.get("is_correct") is True:
            return chr(65 + index)
    return "Not provided"


def validate_assessment_for_docx(questions):
    errors = []
    if not questions:
        raise DocxExportValidationError("questions are required for DOCX export")
    for index, question in enumerate(questions, start=1):
        prefix = f"question {index}"
        if not question.get("metadata"):
            errors.append(f"{prefix}: metadata is required")
        if not str(question.get("body") or "").strip():
            errors.append(f"{prefix}: question body is required")
        if not str(question.get("model_answer") or "").strip():
            errors.append(f"{prefix}: model answer is required")
        if not question.get("quality_checks"):
            errors.append(f"{prefix}: quality-check rows are required")
        if not question.get("revision_options"):
            errors.append(f"{prefix}: revision options are required")

        equations = question.get("equations") or []
        equation_by_label = {
            equation.get("label"): equation
            for equation in equations
            if equation.get("label")
        }
        locations = {
            "question": [question.get("body", "")]
            + [option.get("body", "") for option in question.get("options", [])],
            "solution": [question.get("model_answer", "")],
        }
        for location, texts in locations.items():
            for text in texts:
                for label in _EQUATION_REFERENCE.findall(text or ""):
                    equation = equation_by_label.get(label)
                    if equation is None:
                        errors.append(f"{prefix}: unresolved equation reference {label}")
                    elif equation.get("location") != location:
                        errors.append(
                            f"{prefix}: equation reference {label} belongs in "
                            f"{equation.get('location')}, not {location}"
                        )
    if errors:
        raise DocxExportValidationError("; ".join(errors))


def _solution_units(answer):
    lines = answer.splitlines()
    if len(lines) <= 1:
        lines = re.split(r"(?<=[.!?])\s+(?=[A-Z])", answer)
    for source in lines:
        text = source.strip()
        if not text:
            continue
        text = re.sub(r"^\*\*(.*?)\*\*$", r"\1", text).strip()
        if re.match(r"^Step\s+\d+\s*[:\-\u2013\u2014]", text, re.IGNORECASE):
            continue
        yield text


def _render_solution(document, question):
    rendered_labels = set()
    answer = question["model_answer"]
    segments = question.get("model_answer_segments")
    if segments is not None:
        paragraph = document.add_paragraph(style="Solution Body")
        rendered_labels.update(append_content(
            paragraph, segments, answer,
            equations=question.get("equations", []), location="solution",
        ))
        return rendered_labels

    for unit in _solution_units(answer):
        equation_only = re.fullmatch(r"\[\[EQ:[A-Za-z0-9_-]+\]\][.!]?", unit)
        paragraph = document.add_paragraph(
            style="Solution Equation" if equation_only else "Solution Body"
        )
        if equation_only:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            unit = unit.rstrip(".! ")
        rendered_labels.update(append_content(
            paragraph, None, unit,
            equations=question.get("equations", []), location="solution",
        ))
    return rendered_labels


def _add_answer_key(document, questions):
    mcqs = [
        (index, question)
        for index, question in enumerate(questions, start=1)
        if question.get("type") == "mcq" or question.get("options")
    ]
    if not mcqs:
        return
    document.add_heading("Answer Key", level=2)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    widths = (Inches(2.05), Inches(4.45))
    _set_table_geometry(table, widths)
    table.rows[0].cells[0].text = "Question"
    table.rows[0].cells[1].text = "Correct Answer"
    for cell, width in zip(table.rows[0].cells, widths):
        _format_metadata_cell(cell, width=width, is_label=True)
    for index, question in mcqs:
        cells = table.add_row().cells
        cells[0].text = f"Question {index}: {_question_title(question, index)}"
        cells[1].text = _correct_choice(question)
        _format_metadata_cell(cells[0], width=widths[0], is_label=False)
        _format_metadata_cell(cells[1], width=widths[1], is_label=False)


def _add_quality_check(document, questions):
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    widths = (Inches(1.45), Inches(0.8), Inches(2.45), Inches(0.8), Inches(1.0))
    _set_table_geometry(table, widths)
    headings = ("Criterion", "Rating / 5", "Comment", "User Rating", "User Comment")
    for cell, width, heading in zip(table.rows[0].cells, widths, headings):
        cell.text = heading
        _format_metadata_cell(cell, width=width, is_label=True)
    for index, question in enumerate(questions, start=1):
        for check in question["quality_checks"]:
            cells = table.add_row().cells
            values = (
                f"Q{index}: {check.get('criterion', '')}",
                check.get("rating", ""),
                check.get("comment", ""),
                "",
                "",
            )
            for cell, width, value in zip(cells, widths, values):
                cell.text = str(value)
                _format_metadata_cell(cell, width=width, is_label=False)


def build_assessment_docx(*, run_id: int, prompt_id: int,
                          condition_code: str, run_number: int, course: str, topic: str,
                          questions: list[dict], traceability=None) -> bytes:
    validate_assessment_for_docx(questions)
    document = Document()
    _configure_styles(document)
    document.add_heading("Blueprint Lab Assessment", level=0)
    document.add_heading("1. Assessment Metadata", level=1)
    _add_metadata_table(
        document,
        run_id=run_id,
        prompt_id=prompt_id,
        condition_code=condition_code,
        run_number=run_number,
        course=course,
        topic=topic,
        questions=questions,
        traceability=traceability,
    )

    document.add_heading("2. Student-Facing Questions", level=1)
    for index, question in enumerate(questions, start=1):
        _add_item_heading(document, kind="Question", index=index, question=question)
        paragraph = document.add_paragraph()
        rendered_labels = append_content(
            paragraph,
            question.get("body_segments"),
            question["body"],
            equations=question.get("equations", []),
            location="question",
        )
        for option in question.get("options", []):
            paragraph = document.add_paragraph(style="List Bullet")
            rendered_labels.update(append_content(
                paragraph,
                option.get("segments"),
                option["body"],
                equations=question.get("equations", []),
                location="question",
            ))
        for equation in question.get("equations", []):
            if equation.get("location") != "question":
                continue
            if equation.get("label") in rendered_labels:
                continue
            _add_standalone_equation(document, equation)

    document.add_heading("3. Fully Worked Solution", level=1)
    _add_answer_key(document, questions)
    for index, question in enumerate(questions, start=1):
        _add_item_heading(document, kind="Solution", index=index, question=question)
        rendered_labels = _render_solution(document, question)
        for equation in question.get("equations", []):
            if equation.get("location") != "solution":
                continue
            if equation.get("label") in rendered_labels:
                continue
            _add_standalone_equation(document, equation)

    document.add_heading("4. Assessment Quality Check", level=1)
    _add_quality_check(document, questions)

    document.add_heading("5. Suggested Revision Options", level=1)
    for index, question in enumerate(questions, start=1):
        for revision in question.get("revision_options", []):
            paragraph = document.add_paragraph(style="List Number")
            paragraph.add_run(f"Q{index}: ").bold = True
            paragraph.add_run(revision)

    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()
