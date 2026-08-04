from __future__ import annotations

import hashlib
import io
import zipfile

from docx import Document
from pydantic import ValidationError

from backend.schemas.docx_authoring_schema import RewrittenAssessmentManifest
from backend.services.docx_verification import VerificationIssue, VerificationReport
from backend.services.reproducibility import canonical_json


REQUIRED_HEADINGS = (
    "Assessment Metadata",
    "Questions",
    "Answer Key and Step-by-Step Solutions",
    "Assessment Quality Check",
    "Suggested Revision Options",
)


def _original_question_id(question: dict, ordinal: int) -> str:
    trace = question.get("traceability") or {}
    value = trace.get("assessment_question_id", question.get("id", ordinal + 1))
    return str(value)


class DocxManifestVerifier:
    def verify(self, docx_bytes: bytes, manifest_value: dict, grounding) -> VerificationReport:
        issues: list[VerificationIssue] = []
        manifest_hash = hashlib.sha256(canonical_json(manifest_value).encode()).hexdigest()
        try:
            manifest = RewrittenAssessmentManifest.model_validate(manifest_value)
        except ValidationError as exc:
            issues.append(VerificationIssue("manifest_invalid", evidence=str(exc)[:1000]))
            return VerificationReport(False, tuple(issues), hashlib.sha256(docx_bytes).hexdigest(), manifest_hash)
        originals = grounding.original_assessment.get("questions", [])
        expected = {_original_question_id(question, index) for index, question in enumerate(originals)}
        actual = {question.source_question_id for question in manifest.questions}
        if expected != actual or len(manifest.questions) != len(originals):
            issues.append(VerificationIssue("source_mapping_mismatch", evidence="source mappings do not cover version 1 exactly once"))
        try:
            document = Document(io.BytesIO(docx_bytes))
            visible = "\n".join(
                [paragraph.text for paragraph in document.paragraphs]
                + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
            )
            if len(document.tables) < 3:
                issues.append(VerificationIssue("required_tables_missing", evidence="metadata, answer key, and quality tables are required"))
            if not any(len(table.columns) == 5 for table in document.tables):
                issues.append(VerificationIssue("quality_table_invalid", evidence="quality table must have five columns"))
        except Exception as exc:
            issues.append(VerificationIssue("semantic_mismatch", evidence=f"DOCX text extraction failed: {type(exc).__name__}"))
            return VerificationReport(False, tuple(issues), hashlib.sha256(docx_bytes).hexdigest(), manifest_hash)
        positions = [visible.find(heading) for heading in REQUIRED_HEADINGS]
        if any(position < 0 for position in positions) or positions != sorted(positions):
            issues.append(VerificationIssue("required_sections_missing", evidence="five required headings must exist in order"))
        for label in ("Step", "Final Answer", "Distractor"):
            if label not in visible:
                issues.append(VerificationIssue("solution_labels_missing", evidence=f"{label} label is absent"))
        if "Physical Meaning" not in visible and "Conclusion" not in visible:
            issues.append(VerificationIssue("solution_labels_missing", evidence="interpretation label is absent"))
        for question in manifest.questions:
            required_text = [question.body, question.title]
            required_text.extend(item.body for item in question.options)
            solution = question.solution
            if solution.kind == "computational":
                required_text.extend(solution.knowns_and_target)
                required_text.extend([solution.governing_equation, solution.substitution, *solution.calculation_steps, solution.final_answer, solution.units, solution.physical_meaning])
                required_text.extend(item.explanation for item in solution.distractor_analysis)
            else:
                required_text.extend([solution.governing_concept, *solution.application_steps, solution.conclusion])
                required_text.extend(item.explanation for item in solution.option_elimination)
            key = next(item for item in manifest.answer_key if item.question_id == question.id)
            required_text.extend([key.correct_option_id, key.answer])
            missing = [text for text in required_text if text not in visible]
            if missing:
                issues.append(VerificationIssue("semantic_mismatch", evidence=f"question {question.id} has {len(missing)} manifest values absent from DOCX"))
        for revision in manifest.revision_options:
            if revision not in visible:
                issues.append(VerificationIssue("semantic_mismatch", evidence="revision option absent from DOCX"))
        try:
            with zipfile.ZipFile(io.BytesIO(docx_bytes)) as archive:
                xml = archive.read("word/document.xml")
            if any(q.solution.kind == "computational" for q in manifest.questions) and b"<m:oMath" not in xml:
                issues.append(VerificationIssue("native_math_missing", severity="warning", evidence="computational assessment contains no OOXML math"))
        except Exception:
            pass
        fatal = [item for item in issues if item.severity == "error"]
        return VerificationReport(
            valid=not fatal,
            issues=tuple(issues),
            package_sha256=hashlib.sha256(docx_bytes).hexdigest(),
            manifest_sha256=manifest_hash,
            tool_versions={"manifest_verifier": "1", "python-docx": "1"},
        )
