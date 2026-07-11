from __future__ import annotations

import hashlib
import json
from importlib.metadata import version as package_version
from io import BytesIO
from zipfile import BadZipFile

from docx import Document
from pypdf import PdfReader
from pypdf.errors import PdfReadError
from sqlalchemy.orm import Session

from backend.models.source_document import SourceDocument

MAX_SOURCE_DOCUMENT_BYTES = 20 * 1024 * 1024


class SourceDocumentValidationError(ValueError):
    def __init__(self, code: str):
        self.code = code
        super().__init__(code)


def _plain(content: bytes):
    try:
        return content.decode("utf-8"), "plain-text:utf-8"
    except UnicodeDecodeError as exc:
        raise SourceDocumentValidationError("invalid_source_document") from exc


def _json(content: bytes):
    text, _ = _plain(content)
    try:
        json.loads(text)
    except json.JSONDecodeError as exc:
        raise SourceDocumentValidationError("invalid_source_document") from exc
    return text, "plain-text:utf-8"


def _docx(content: bytes):
    try:
        text = "\n".join(p.text for p in Document(BytesIO(content)).paragraphs if p.text.strip())
    except Exception as exc:
        raise SourceDocumentValidationError("invalid_source_document") from exc
    return text, f"python-docx:{package_version('python-docx')}"


def _pdf(content: bytes):
    try:
        reader = PdfReader(BytesIO(content), strict=True)
        if reader.is_encrypted:
            raise SourceDocumentValidationError("encrypted_source_document")
        text = "\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()
    except SourceDocumentValidationError:
        raise
    except Exception as exc:
        raise SourceDocumentValidationError("invalid_source_document") from exc
    return text, f"pypdf:{package_version('pypdf')}"


EXTRACTORS = {
    "text/plain": _plain,
    "text/markdown": _plain,
    "application/json": _json,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": _docx,
    "application/pdf": _pdf,
}


def create_source_document(db: Session, *, name: str, document_type: str, version: str,
        filename: str, media_type: str, content: bytes, description: str | None) -> SourceDocument:
    if len(content) > MAX_SOURCE_DOCUMENT_BYTES:
        raise SourceDocumentValidationError("source_document_too_large")
    extractor = EXTRACTORS.get(media_type)
    if extractor is None:
        raise SourceDocumentValidationError("unsupported_source_document_media_type")
    extracted_text, extraction_method = extractor(content)
    if not extracted_text.strip():
        raise SourceDocumentValidationError("empty_source_document")
    source = SourceDocument(name=name, document_type=document_type, version=version,
        original_filename=filename, media_type=media_type, content=content,
        content_hash=hashlib.sha256(content).hexdigest(), extracted_text=extracted_text,
        extraction_method=extraction_method, description=description or "")
    db.add(source)
    try:
        db.commit()
        db.refresh(source)
    except Exception:
        db.rollback()
        raise
    return source
